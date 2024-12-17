from flask import Flask, request, render_template, send_file
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import os
import pickle
from datetime import datetime
from docx import Document

app = Flask(__name__)

# Escopos de acesso à API
SCOPES = ['https://www.googleapis.com/auth/tasks.readonly']

def authenticate_google_tasks():
    """Autentica e retorna o serviço da API do Google Tasks."""
    creds = None
    # Carregar token salvo
    if os.path.exists('token.pkl'):
        with open('token.pkl', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pkl', 'wb') as token:
            pickle.dump(creds, token)
    return build('tasks', 'v1', credentials=creds)

def get_completed_tasks(service, start_date, end_date):
    """Obtém as tarefas concluídas dentro de um intervalo de datas."""
    start_date_iso = start_date.isoformat() + 'Z'
    end_date_iso = end_date.isoformat() + 'Z'
    completed_tasks = []

    # Obter listas de tarefas
    tasklists = service.tasklists().list().execute()
    for tasklist in tasklists.get('items', []):
        tasks = service.tasks().list(
            tasklist=tasklist['id'],
            showCompleted=True,
            showHidden=True,
            completedMin=start_date_iso,
            completedMax=end_date_iso
        ).execute()
        for task in tasks.get('items', []):
            if task.get('status') == 'completed':
                completed_tasks.append({
                    'title': task.get('title', 'Sem Título'),
                    'completed': task.get('completed')
                })
    return completed_tasks

def generate_word_report(tasks, start_date, end_date):
    """Gera um relatório em Word das tarefas concluídas."""
    output_file = "relatorio_tarefas.docx"
    document = Document()
    document.add_heading('Relatório de Tarefas Concluídas', level=1)
    document.add_paragraph(f"Período: {start_date.strftime('%d/%m/%Y')} a {end_date.strftime('%d/%m/%Y')}\n")

    if not tasks:
        document.add_paragraph('Nenhuma tarefa concluída no período especificado.')
    else:
        for task in tasks:
            completed_date = datetime.fromisoformat(task['completed'].replace('Z', ''))
            document.add_paragraph(f"- {task['title']} (Concluída em: {completed_date.strftime('%d/%m/%Y %H:%M')})")

    document.save(output_file)
    return output_file

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            # Obter datas do formulário
            start_date = datetime.strptime(request.form["start_date"], "%Y-%m-%d")
            end_date = datetime.strptime(request.form["end_date"], "%Y-%m-%d")

            # Autenticação e obtenção de tarefas
            service = authenticate_google_tasks()
            tasks = get_completed_tasks(service, start_date, end_date)

            # Gerar relatório
            report_file = generate_word_report(tasks, start_date, end_date)
            return send_file(report_file, as_attachment=True)
        except Exception as e:
            return f"Erro: {e}"

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)

